"""Text extraction from .doc and .docx files."""

import struct
import zipfile
from pathlib import Path

import olefile
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


# Maximum file size for .doc files (50 MB)
_MAX_DOC_SIZE_BYTES = 50 * 1024 * 1024

# Error message prefix
_ERROR_PREFIX = "Error: Extraction - "


class DocxExtractor:
    """Extract text from .docx (XML/ZIP) files using python-docx."""

    def extract(self, file_path: Path) -> str:
        """Extract text from a .docx file.

        Returns extracted text joined by newlines, or an error message string
        starting with 'Error: ' on failure, or empty string if no text content.
        """
        try:
            doc = Document(str(file_path))
        except zipfile.BadZipFile:
            return f"{_ERROR_PREFIX}文件不是有效的 .docx 文档"
        except PackageNotFoundError:
            return f"{_ERROR_PREFIX}文件不是有效的 .docx 文档"
        except Exception as e:
            # Detect encrypted .docx files
            # python-docx may raise various errors for encrypted files
            error_msg = str(e).lower()
            if "encrypt" in error_msg or "password" in error_msg:
                return f"{_ERROR_PREFIX}文件已加密，无法读取"
            # Check if it's a corrupted/invalid file
            return f"{_ERROR_PREFIX}文件不是有效的 .docx 文档"

        # Check for encryption by examining the file as a ZIP
        # Encrypted .docx files may open but contain EncryptedPackage
        try:
            if self._is_encrypted_docx(file_path):
                return f"{_ERROR_PREFIX}文件已加密，无法读取"
        except Exception:
            pass

        # Extract paragraphs text
        texts: list[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                texts.append(paragraph.text)

        # Extract table cells text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        texts.append(cell.text)

        if not texts:
            return ""

        return "\n".join(texts)

    def _is_encrypted_docx(self, file_path: Path) -> bool:
        """Check if a .docx file is actually an encrypted OLE container."""
        # Some encrypted .docx files are actually OLE files with EncryptedPackage
        if olefile.isOleFile(str(file_path)):
            with olefile.OleFileIO(str(file_path)) as ole:
                if ole.exists("EncryptedPackage"):
                    return True
        return False


class DocExtractor:
    """Extract text from .doc (binary OLE) files using olefile."""

    def extract(self, file_path: Path) -> str:
        """Extract text from a .doc file.

        Returns extracted text as UTF-8 string, or an error message string
        starting with 'Error: ' on failure, or empty string if no text content.
        """
        # Check file size limit (50 MB)
        try:
            file_size = file_path.stat().st_size
        except OSError:
            return f"{_ERROR_PREFIX}无法访问文件"

        if file_size > _MAX_DOC_SIZE_BYTES:
            return f"{_ERROR_PREFIX}文件超出支持的最大大小（50 MB）"

        # Validate it's a valid OLE file
        if not olefile.isOleFile(str(file_path)):
            return f"{_ERROR_PREFIX}文件不是有效的 .doc 文档"

        try:
            with olefile.OleFileIO(str(file_path)) as ole:
                # Check for encryption
                if self._is_encrypted(ole):
                    return f"{_ERROR_PREFIX}文件受保护，无法提取"

                # Read the WordDocument stream
                if not ole.exists("WordDocument"):
                    return f"{_ERROR_PREFIX}文件不是有效的 .doc 文档"

                word_stream = ole.openstream("WordDocument").read()
                return self._extract_text_from_word_stream(ole, word_stream)

        except Exception as e:
            error_msg = str(e).lower()
            if "encrypt" in error_msg or "password" in error_msg:
                return f"{_ERROR_PREFIX}文件受保护，无法提取"
            return f"{_ERROR_PREFIX}文件不是有效的 .doc 文档"

    def _is_encrypted(self, ole: olefile.OleFileIO) -> bool:
        """Check if the .doc file is encrypted.

        Word documents use FIB (File Information Block) flags to indicate encryption.
        Bit 0x0100 of FIB.fib_base.flags indicates encryption.
        Also checks for the presence of encryption-related streams.
        """
        # Check for EncryptedPackage stream (OOXML encryption)
        if ole.exists("EncryptedPackage"):
            return True

        # Check FIB flags in WordDocument stream
        try:
            if ole.exists("WordDocument"):
                word_stream = ole.openstream("WordDocument").read()
                if len(word_stream) >= 12:
                    # FIB base flags are at offset 10 (bytes 10-11)
                    flags = struct.unpack_from("<H", word_stream, 10)[0]
                    # Bit 8 (0x0100) indicates encryption
                    if flags & 0x0100:
                        return True
        except Exception:
            pass

        return False

    def _extract_text_from_word_stream(
        self, ole: olefile.OleFileIO, word_stream: bytes
    ) -> str:
        """Extract text content from the Word binary format.

        Uses the FIB (File Information Block) to locate text in the
        CLX (Complex) part of the file, or falls back to reading from
        the Table stream.
        """
        if len(word_stream) < 12:
            return ""

        # Read FIB base to determine byte order and version
        # wIdent should be 0xA5EC for Word documents
        w_ident = struct.unpack_from("<H", word_stream, 0)[0]
        if w_ident != 0xA5EC:
            return ""

        # Get FIB flags
        flags = struct.unpack_from("<H", word_stream, 10)[0]

        # Check if the document uses complex format
        # Bit 2 (0x0004) = fComplex
        is_complex = bool(flags & 0x0004)

        # Get ccpText - character count of main document text
        # Located at FIB offset for fcMin/ccpText area
        # In FIB, the character counts are in FibRgLw section
        # FibBase is 32 bytes, then csw(2) + fibRgW variable, then cslw(2) + FibRgLw
        # Simplified: try to extract text from the document streams

        try:
            text = self._extract_text_simple(ole, word_stream)
            return text
        except Exception:
            return ""

    def _extract_text_simple(
        self, ole: olefile.OleFileIO, word_stream: bytes
    ) -> str:
        """Simple text extraction from Word binary document.

        Reads the text from the clx/piece table in the table stream,
        or extracts directly from the WordDocument stream.
        """
        if len(word_stream) < 24:
            return ""

        # Read ccpText from FibRgLw97 structure
        # FibBase = 32 bytes
        # Then comes: csw (2 bytes) + FibRgW97 (28 bytes = 14 uint16)
        # Then comes: cslw (2 bytes) + FibRgLw97
        # ccpText is the first field of FibRgLw97

        fib_base_size = 32
        if len(word_stream) < fib_base_size + 2:
            return ""

        csw = struct.unpack_from("<H", word_stream, fib_base_size)[0]
        fibrg_w_offset = fib_base_size + 2
        fibrg_w_size = csw * 2  # csw is count of uint16

        cslw_offset = fibrg_w_offset + fibrg_w_size
        if len(word_stream) < cslw_offset + 2:
            return ""

        cslw = struct.unpack_from("<H", word_stream, cslw_offset)[0]
        fibrg_lw_offset = cslw_offset + 2

        if len(word_stream) < fibrg_lw_offset + 4:
            return ""

        # ccpText is the first uint32 in FibRgLw97
        ccp_text = struct.unpack_from("<I", word_stream, fibrg_lw_offset)[0]

        if ccp_text == 0:
            return ""

        # Determine which table stream to use (0Table or 1Table)
        flags = struct.unpack_from("<H", word_stream, 10)[0]
        # Bit 9 (0x0200) = fWhichTblStm: 0 = 0Table, 1 = 1Table
        table_stream_name = "1Table" if (flags & 0x0200) else "0Table"

        if not ole.exists(table_stream_name):
            # Fallback: try to read text directly from WordDocument stream
            return self._extract_text_direct(word_stream, ccp_text)

        table_stream = ole.openstream(table_stream_name).read()

        # Get CLX offset and size from FibRgFcLcb
        # FibRgFcLcb starts after FibRgLw
        fibrg_lw_size = cslw * 4
        fcclcb_offset = fibrg_lw_offset + fibrg_lw_size

        if len(word_stream) < fcclcb_offset + 2:
            return self._extract_text_direct(word_stream, ccp_text)

        # cbRgFcLcb count
        cbrgfclcb = struct.unpack_from("<H", word_stream, fcclcb_offset)[0]
        fcclcb_data_offset = fcclcb_offset + 2

        # fcClx is at index 66 in FibRgFcLcb97 (each entry is fc:4 + lcb:4 = 8 bytes)
        # fcClx offset = fcclcb_data_offset + 66 * 8
        clx_index = 66
        clx_fc_offset = fcclcb_data_offset + clx_index * 8

        if len(word_stream) < clx_fc_offset + 8:
            return self._extract_text_direct(word_stream, ccp_text)

        fc_clx = struct.unpack_from("<I", word_stream, clx_fc_offset)[0]
        lcb_clx = struct.unpack_from("<I", word_stream, clx_fc_offset + 4)[0]

        if fc_clx == 0 or lcb_clx == 0:
            return self._extract_text_direct(word_stream, ccp_text)

        if fc_clx + lcb_clx > len(table_stream):
            return self._extract_text_direct(word_stream, ccp_text)

        # Parse the CLX structure to get piece table
        clx_data = table_stream[fc_clx : fc_clx + lcb_clx]
        return self._parse_piece_table(clx_data, word_stream, ccp_text)

    def _parse_piece_table(
        self, clx_data: bytes, word_stream: bytes, ccp_text: int
    ) -> str:
        """Parse the CLX piece table to extract text."""
        offset = 0
        text_parts: list[str] = []

        # Skip any Prc (prefix) entries (type 0x01)
        while offset < len(clx_data) and clx_data[offset] == 0x01:
            if offset + 3 > len(clx_data):
                break
            cb_grpprl = struct.unpack_from("<H", clx_data, offset + 1)[0]
            offset += 3 + cb_grpprl

        # Now we should be at the Pcdt (type 0x02)
        if offset >= len(clx_data) or clx_data[offset] != 0x02:
            return self._extract_text_direct(word_stream, ccp_text)

        offset += 1  # skip type byte

        if offset + 4 > len(clx_data):
            return self._extract_text_direct(word_stream, ccp_text)

        lcb = struct.unpack_from("<I", clx_data, offset)[0]
        offset += 4

        # PlcPcd structure: array of CPs followed by array of PCD entries
        # Number of pieces = (lcb - 4) / (4 + 8) ... approximately
        # CPs array has n+1 entries (each 4 bytes), PCDs array has n entries (each 8 bytes)
        # lcb = (n+1)*4 + n*8 = 4n + 4 + 8n = 12n + 4
        # n = (lcb - 4) / 12

        if lcb < 4:
            return ""

        n = (lcb - 4) // 12
        if n <= 0:
            return self._extract_text_direct(word_stream, ccp_text)

        pcd_offset_base = offset

        # Read character positions (CPs)
        cps: list[int] = []
        for i in range(n + 1):
            cp_off = pcd_offset_base + i * 4
            if cp_off + 4 > len(clx_data):
                break
            cps.append(struct.unpack_from("<I", clx_data, cp_off)[0])

        # Read piece descriptors
        pcd_start = pcd_offset_base + (n + 1) * 4
        for i in range(min(n, len(cps) - 1)):
            pcd_off = pcd_start + i * 8
            if pcd_off + 8 > len(clx_data):
                break

            # PCD structure: 2 bytes descriptor, 4 bytes fc, 2 bytes prm
            # fc field encodes the file offset and whether it's compressed
            fc_value = struct.unpack_from("<I", clx_data, pcd_off + 2)[0]

            # Bit 30 indicates compressed (ANSI) text
            is_compressed = bool(fc_value & 0x40000000)
            # Clear the compressed bit to get the actual offset
            fc_offset = fc_value & 0x3FFFFFFF

            cp_start = cps[i]
            cp_end = cps[i + 1]
            char_count = cp_end - cp_start

            if char_count <= 0:
                continue

            if is_compressed:
                # Compressed: 1 byte per character (cp1252/latin-1)
                byte_offset = fc_offset // 2
                byte_count = char_count
                if byte_offset + byte_count <= len(word_stream):
                    raw = word_stream[byte_offset : byte_offset + byte_count]
                    text_parts.append(raw.decode("cp1252", errors="replace"))
            else:
                # Uncompressed: 2 bytes per character (UTF-16LE)
                byte_count = char_count * 2
                if fc_offset + byte_count <= len(word_stream):
                    raw = word_stream[fc_offset : fc_offset + byte_count]
                    text_parts.append(raw.decode("utf-16-le", errors="replace"))

        result = "".join(text_parts)

        # Only take ccp_text characters (main document body)
        if len(result) > ccp_text:
            result = result[:ccp_text]

        # Clean up: replace \r with \n for paragraph breaks
        result = result.replace("\r\n", "\n").replace("\r", "\n")

        # Remove null characters and other control characters except \n and \t
        cleaned = []
        for ch in result:
            if ch == "\n" or ch == "\t" or (ch >= " "):
                cleaned.append(ch)
        result = "".join(cleaned)

        return result.strip()

    def _extract_text_direct(self, word_stream: bytes, ccp_text: int) -> str:
        """Fallback: try to extract text directly by scanning for readable content."""
        # This is a simplified fallback that tries to find text patterns
        # in the word stream when piece table parsing fails
        if len(word_stream) < 512:
            return ""

        # Try to find and decode text content after the FIB header
        # Word documents typically store text starting after a header area
        # Try decoding as UTF-16LE from various offsets
        best_text = ""
        # Scan for UTF-16LE text patterns
        for start_offset in range(512, min(len(word_stream), 4096), 2):
            try:
                chunk = word_stream[start_offset : start_offset + min(ccp_text * 2, 8192)]
                text = chunk.decode("utf-16-le", errors="ignore")
                # Check if it looks like actual text
                printable_ratio = sum(
                    1 for c in text if c.isprintable() or c in "\n\r\t"
                ) / max(len(text), 1)
                if printable_ratio > 0.7 and len(text) > len(best_text):
                    best_text = text
            except Exception:
                continue

        if best_text:
            # Clean up
            best_text = best_text.replace("\r\n", "\n").replace("\r", "\n")
            cleaned = []
            for ch in best_text:
                if ch == "\n" or ch == "\t" or (ch >= " "):
                    cleaned.append(ch)
            return "".join(cleaned).strip()

        return ""


class TextExtractor:
    """Main text extractor that routes to format-specific extractors."""

    def __init__(self) -> None:
        self._docx_extractor = DocxExtractor()
        self._doc_extractor = DocExtractor()

    def extract(self, file_path: Path) -> str:
        """Extract text from .doc or .docx file.

        Routes to the appropriate extractor based on file extension (case-insensitive).
        Returns extracted text, empty string for no content, or error message
        starting with 'Error: ' on failure.
        """
        extension = file_path.suffix.lower()

        if extension == ".docx":
            return self._docx_extractor.extract(file_path)
        elif extension == ".doc":
            return self._doc_extractor.extract(file_path)
        else:
            return f"{_ERROR_PREFIX}不支持的文件格式: {extension}"
